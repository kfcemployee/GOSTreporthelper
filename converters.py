from style_configs.style_config import StyleNames, StyleManager
from docx.document import Document
from docx import Document as Doc
import re
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

class Converter:
    def __init__(self, doc: Document, data: list, output_path: str):
        # главные параметры: объект документа, данные парсинга и место для сохранения
        self.doc = doc
        self.data = data
        self.output_path = output_path

        # объект класса стилей
        self.style_conf = StyleManager()
        self.style_conf.setup_styles(self.doc)

        # переменные для сквозной нумерации (для таблиц не надо)
        self.img_counter = 1
        self.formula_counter = 1

    def format_pages(self):
        for i, page in enumerate(self.doc.sections):
            page.page_height = Cm(29.7)
            page.page_width = Cm(21)
            page.orientation = WD_ORIENTATION.PORTRAIT
            page.top_margin = Cm(2)
            page.bottom_margin = Cm(2)
            page.left_margin = Cm(3)
            page.right_margin = Cm(1.5)

            if i > 0:
                p = page.footer.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.style = StyleNames.normal
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(0)

                run = p.add_run()

                fld_char_begin = OxmlElement('w:fldChar')
                fld_char_begin.set(qn('w:fldCharType'), 'begin')

                instr_text = OxmlElement('w:instrText')
                instr_text.set(qn('xml:space'), 'preserve')
                instr_text.text = 'PAGE'

                fld_char_end = OxmlElement('w:fldChar')
                fld_char_end.set(qn('w:fldCharType'), 'end')

                run._element.append(fld_char_begin)
                run._element.append(instr_text)
                run._element.append(fld_char_end)

    # форматирование таблиц
    def format_tables(self):
        for i, tb in enumerate(self.doc.tables, 1):
            elem = tb._element
            tb.alignment = WD_TABLE_ALIGNMENT.CENTER

            parent = elem.getparent()
            prev, pnext = elem.getprevious(), elem.getnext()

            if prev is not None and (prev.tag.endswith('p') and "таблица" in
                     prev.text.lower()):
                text = prev.text

                tname = re.findall(r"^Таблица[\s:\-−.]?(.+)", text)
                if tname:
                    fcap = self.doc.add_paragraph(text=f"Таблица {i} − {tname[0].strip()}")
                else:
                    fcap = self.doc.add_paragraph(text=f"Таблица {i} − []")

                fcap.style = StyleNames.caption
                parent.insert(parent.index(prev), fcap._element)
                parent.remove(prev)
            elif pnext is not None and (pnext.tag.endswith('p') and "таблица" in
                    pnext.text.lower()):
                text = pnext.text

                tname = re.findall(r"^Таблица[\s:\-.]?(.+)", text)
                if tname:
                    fcap = self.doc.add_paragraph(text=f"Таблица {i} − {tname[0].strip()}")
                else:
                    fcap = self.doc.add_paragraph(text=f"Таблица {i} − []")

                fcap.style = StyleNames.caption
                parent.insert(parent.index(elem), fcap._element)
                parent.remove(pnext)
            else:
                fcap = self.doc.add_paragraph()
                fcap.style = StyleNames.caption
                fcap.add_run(f"Таблица {i} − []")
                parent.insert(parent.index(elem), fcap._element)

            fcap.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            tbl_pr = tb._element.tblPr

            borders_xml = f'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                borders_xml += (f'<w:{border_name} w:val="single" w:sz="{self.style_conf.tb_conf.get("border_size")}" '
                                f'w:space="0" w:color="000000"/>')
            borders_xml += '</w:tblBorders>'

            for b in tbl_pr.xpath('.//w:tblBorders'):
                b.getparent().remove(b)

            tbl_pr.append(parse_xml(borders_xml))

            for ix, row in enumerate(tb.rows):
                row.height = self.style_conf.tb_conf.get("row_height")
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

                for cell in row.cells:
                    tc_pr = cell._element.tcPr
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    mar = tc_pr.find(qn("w:tcMar"))
                    if mar is not None:
                        tc_pr.remove(mar)

                    mar = OxmlElement("w:tcMar")
                    for side in ['left', 'right']:
                        mr_e = OxmlElement(f"w:{side}")
                        mr_e.set(qn("w:w"), "100")
                        mr_e.set(qn("w:type"), "dxa")

                        mar.append(mr_e)
                    tc_pr.append(mar)

                    for p in cell.paragraphs:
                        p.style = StyleNames.tb_item
                        if ix == 0:
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for r in p.runs:
                                r.bold = True
                        else:
                            if re.match(r"\d+", p.text):
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                            else:
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def format_doc(self, c):

        if "heading" in c["ptype"]:
            self.format_headings(c["par"], c["level"])
        if c["ptype"] == 'list_bullet':
            print(c["par"].text)
            self.format_bullet(c["par"])
        if c["ptype"] == "list_number":
            print(c["par"].text)
            self.format_numbered(c["par"])
        if c["ptype"] == "normal":
            self.format_normal(c["par"])
        if c["ptype"] == 'image':
            self.format_image(c["par"])
        if c["ptype"] == "code":
            self.format_code(c["par"])
        if c["ptype"] == "empty":
            self.format_empty(c["par"], c["index"])

    @staticmethod
    def format_code(par: Paragraph):
        par.style = StyleNames.code

    def format_formula(self, par: Paragraph):
        par.style = StyleNames.formula
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        par.text = f"{re.sub(r'\s*\(\d+\)\s*$', '', par.text.strip())},                    ({self.formula_counter})"
        for run in par.runs:
            run.font.italic = True

    @staticmethod
    def format_headings(par, level):
        res = re.findall(r"([\d.]+\.)\s(.+)", par.text)
        dig, text = None, par.text
        if res:
            dig, text = res[0]
        match level:
            case 1:
                par.style = StyleNames.h1
                par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            case 2:
                par.style = StyleNames.h2
                par.text = text.capitalize()
            case 3:
                par.style = StyleNames.h3
                par.text = text.capitalize()
        if dig:
            if dig.endswith('.'):
                par.text = dig[:-1] + ' ' + text.capitalize()

    @staticmethod
    def format_bullet(par: Paragraph):
        p_elem = par._element
        p_pr = p_elem.get_or_add_pPr()
        for elem_name in ['numPr', 'numId', 'ilvl', 'numbering']:
            elem = p_pr.find(qn(f'w:{elem_name}'))
            if elem is not None:
                p_pr.remove(elem)
        if p_pr.numPr is not None:
            p_pr.remove(p_pr.numPr)
        par.style = StyleNames.list
        par.paragraph_format.first_line_indent = Cm(0)
        par.paragraph_format.left_indent = Cm(1.25)

        text = par.text.strip()
        if text and not text.startswith('− '):
            for run in par.runs:
                run.text = ""
            par.add_run('− ' + re.sub(r"^[*•\-]\s?", "", text))

    def format_empty(self, p, index):
        if ('<w:drawing' in p._element.xml) or ('<w:pict' in p._element.xml):
            self.format_image(p)
            return

        if "w:br" in p._element.xml:
            return

        if not self.keep_empty(index, self.data):
            p_element = p._element
            p_element.getparent().remove(p_element)

    @staticmethod
    def keep_empty(index: int, allp):
        if index in [0, len(allp) - 1]:
            return False

        if (((None, allp[index - 1]['ptype'])[index > 0], (None, allp[index + 1]['ptype'])[index < len(allp) - 1]) in
                [('header', 'normal'),
                 ('header', 'header'),
                 ('caption', 'normal')]):
            return True
        return False

    def format_numbered(self, par: Paragraph):
        par.style = StyleNames.list

        if not re.match(r'^[\dа-яa-z]+[).]\s+.+', par.text, re.IGNORECASE):
            self.format_bullet(par)

    @staticmethod
    def format_normal(par: Paragraph):
        par.style = StyleNames.normal
        par.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def format_image(self, par: Paragraph):
        elem = par._element

        prev, pnext = elem.getprevious(), elem.getnext()
        parent = elem.getparent()

        if prev is not None and (prev.tag.endswith('p') and
                re.match(r"^Рисунок|Рис\.\s.+", prev.text.lower())):
            text = prev.text

            tname = re.findall(r"^(?:Рисунок|Рис\.)[\s:\-.]?(.+)", text)
            if tname:
                fcap = self.doc.add_paragraph(text=f"Таблица {self.img_counter} − {tname[0].strip()}")
            else:
                fcap = self.doc.add_paragraph(text=f"Таблица {self.img_counter} − []")
            parent.insert(parent.index(elem) + 1, fcap._element)
            parent.remove(prev)
        elif pnext is not None and (pnext.tag.endswith('p') and
              re.match(r"^Рисунок|Рис\.\s.+", pnext.text.lower())):
            text = pnext.text

            tname = re.findall(r"^(?:Рисунок|Рис\.)[\s:\-.]?(.+)", text)
            if tname:
                fcap = self.doc.add_paragraph(text=f"Рисунок {self.img_counter} − {tname[0].strip()}")
            else:
                fcap = self.doc.add_paragraph(text=f"Рисунок {self.img_counter} − []")
            parent.insert(parent.index(elem) + 1, fcap._element)
            parent.remove(pnext)
        else:
            fcap = self.doc.add_paragraph()
            fcap.add_run(f"Рисунок {self.img_counter} − []")
            parent.insert(parent.index(elem) + 1, fcap.element)

        fcap.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fcap.style = StyleNames.caption
        par.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        par.paragraph_format.space_before = Pt(12)
        par.paragraph_format.space_after = Pt(6)

        self.img_counter += 1

    def start(self):
        self.format_pages()
        self.format_tables()

        for c in self.data:
            self.format_doc(c)

        self.doc.save(self.output_path)

class MarkdownConverter:
    def __init__(self, data: list, output_path: str):
        self.data = data
        self.doc = Doc()
        self.output_path = output_path

    def convert_to_doc(self):
        last_p, tb = {}, None

        index = -1
        for p in self.data:
            cur_el = {}

            if "table" not in p["type"]:
                par = self.doc.add_paragraph()

                if p["type"] == 'page_break':
                    par.paragraph_format.page_break_before = True

                if p["type"] == "header":
                    par.style = f"Heading {min(3, int(p["level"]))}"
                    cur_el["par"], cur_el["text"] = par, p["text"].replace('**', '').strip()
                    cur_el["ptype"] = "header"
                    par.add_run(p["text"].replace('**', '').strip())

                if p["type"] == 'normal':
                    par.style = "Normal"
                    cur_el["text"] = p["text"]
                    cur_el["ptype"] = "normal"
                    self._add_run_f(par, p["text"])

                if p["type"] in ['ord_list', 'unord_list']:
                    par.style = "List Number"
                    cur_el["text"] = p["text"]
                    cur_el["ptype"] = "list_item"
                    if p["type"] == "unord_list":
                        par.style = "List Bullet"
                        par.add_run("— ")

                    self._add_run_f(par, p["text"])

                if p["type"] == "empty":
                    cur_el["par"], cur_el["text"] = par, p
                    cur_el["ptype"] = "empty"

                index += 1
            else:
                col_text = [i.strip() for i in p["text"].split('|') if i]
                if "table" in last_p["type"]:
                    if p["type"] == "table":
                        c_row = tb.add_row()
                        for k, text in enumerate(col_text):
                            if text:
                                if re.fullmatch(r'[\d., ]+', text):
                                    (c_row.cells[k].paragraphs[0]
                                     .paragraph_format.alignment) = WD_PARAGRAPH_ALIGNMENT.RIGHT
                                else:
                                    c_row.cells[k].paragraphs[
                                        0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                self._add_run_f(c_row.cells[k].paragraphs[0], text)
                else:
                    tb = self.doc.add_table(rows=1, cols=len(col_text))
                    for k, i in enumerate(col_text):
                        if i:
                            tb.rows[0].cells[k].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            self._add_run_f(tb.rows[0].cells[k].paragraphs[0], i)

            last_p = p
        self.doc.save(self.output_path)
        return self.doc

    @staticmethod
    def detect_inline_format( text):
        b_pat, ind = r"\*\*(.*?)\*\*", []
        italic_pat = r'(?<!\*)\*(?!\*)([^*]+?)(?<!\*)\*(?!\*)'

        for match in re.finditer(italic_pat, text):
            ind += [[match.span(), 'italic']]

        for match in re.finditer(b_pat, text):
            ind += [[match.span(), 'bold']]

        if ind:
            ind.sort(key=lambda x: x[0][0])
            bounds = [[(0, ind[0][0][0]), 'normal']]
            for i in range(len(ind)):
                bounds += [ind[i]]

                if i < len(ind) - 1:
                    bounds += [[(ind[i][0][1], ind[i + 1][0][0]), 'normal']]
                else:
                    bounds += [[(ind[i][0][1], len(text)), 'normal']]
        else:
            bounds = [[(0, len(text)), 'normal']]
        return bounds

    def _add_run_f(self, pg, text):
        bounds = self.detect_inline_format(text)
        bounds.sort(key=lambda x: x[0][0])

        for i in bounds:
            (st, e), style = i
            cur_run = pg.add_run(text[st:e].replace("*", ""))
            cur_run.font.size = Pt(14)
            cur_run.font.name = "Times New Roman"

            match style:
                case "bold":
                    cur_run.bold = True
                case "italic":
                    cur_run.italic = True
